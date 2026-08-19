import sys
import json
import docx
import copy
from datetime import datetime

# Helper to format numbers with spaces as thousands separator
def format_money(val):
    return "{:,.2f}".format(val).replace(",", "\xa0")

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    else:
        tcMar.clear()
        
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

def remove_cell_horizontal_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # Set top and bottom borders to nil
    for border_name in ['top', 'bottom']:
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is not None:
            tcBorders.remove(border)
        new_border = OxmlElement(f'w:{border_name}')
        new_border.set(qn('w:val'), 'nil')
        tcBorders.append(new_border)

def set_cell_text(cell, text, right_align=False, font_size=12):
    # The cell text is in the last paragraph
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[-1]
    
    # Reset paragraph formatting and indents
    p.paragraph_format.left_indent = docx.shared.Pt(0)
    p.paragraph_format.right_indent = docx.shared.Pt(0)
    p.paragraph_format.space_before = docx.shared.Pt(2)
    p.paragraph_format.space_after = docx.shared.Pt(2)
    p.paragraph_format.line_spacing = 1.0
    
    if right_align:
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        
    # Clear extra runs
    for r in p.runs[1:]:
        p._p.remove(r._r)
        
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
        
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(font_size)
    run.font.color.rgb = docx.shared.RGBColor(0, 0x1F, 0x5F)

def main():
    if len(sys.argv) < 3:
        print("Usage: python generate-devis-docx.py <json_path> <output_path>")
        sys.exit(1)
        
    json_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Load template docx
    template_path = "C:\\Users\\othaad\\Desktop\\OXY-CHIMI\\CHMIRAL\\CHMIRAL\\backend\\assets\\DEVIS SINOSPARKS FLINKOTE CH-064-25-0.docx"
    doc = docx.Document(template_path)

    # Find paragraphs dynamically
    p1 = None
    p9 = None
    p14 = None
    p35 = doc.paragraphs[35] if len(doc.paragraphs) > 35 else None
    p45 = None

    for p in doc.paragraphs:
        txt = p.text.strip()
        if "DEVIS N" in txt and "A rappeler" not in txt:
            p1 = p
        elif txt.startswith("Casablanca le"):
            p9 = p
        elif txt.startswith("Objet :"):
            p14 = p
        elif txt.startswith("Préconisations"):
            p45 = p

    middle_num = data.get("numeroDevisMiddle", "064")

    # 1. Update Paragraph 1 (DEVIS N° : CH-XXX-26)
    if p1:
        if len(p1.runs) > 8:
            p1.runs[8].text = f"{middle_num}-"
        else:
            p1.text = f"DEVIS N° : CH-{middle_num}-26"

    # 2. Update Paragraph 9 (Casablanca le : DD/MM/YYYY)
    if p9:
        date_str = data.get("dateDevis", "")
        if date_str:
            try:
                dt = datetime.strptime(date_str, "%Y-%m-%d")
                formatted_date = dt.strftime("%d/%m/%Y")
            except Exception:
                formatted_date = date_str
        else:
            formatted_date = datetime.now().strftime("%d/%m/%Y")
            
        if len(p9.runs) > 4:
            p9.runs[4].text = f":{formatted_date}"
        else:
            p9.text = f"Casablanca le :{formatted_date}"

    # 3. Update Paragraph 14 (Objet : FLINKOTE.)
    if p14:
        objet_text = data.get("objet", "FLINKOTE")
        if not objet_text.endswith('.'):
            objet_text += '.'
            
        if len(p14.runs) > 4:
            run = p14.runs[4]
            run.text = objet_text
            run.bold = True
            run.font.size = docx.shared.Pt(15)
            run.font.name = 'Times New Roman'
        else:
            p14.text = ""
            r0 = p14.add_run("Objet : ")
            r0.bold = True
            r0.font.name = 'Times New Roman'
            r0.font.size = docx.shared.Pt(16)
            
            r1 = p14.add_run(objet_text)
            r1.bold = True
            r1.font.name = 'Times New Roman'
            r1.font.size = docx.shared.Pt(15)

    # 4. Clear Paragraph 35 text runs so it remains blank but keeps its exact line height and layout.
    if p35:
        for r in p35.runs:
            r.text = ""

    # 5. Update Paragraph 45 (Préconisations : réf Devis CH-XXX-26) and remove 2 empty spaces before it
    if p45:
        if len(p45.runs) > 8:
            p45.runs[8].text = f"CH-{middle_num}-"
        else:
            p45.text = f"Préconisations :\tréf Devis CH-{middle_num}-26"
            
        # Delete two empty paragraphs right before p45
        try:
            p45_idx = None
            for idx, p in enumerate(doc.paragraphs):
                if p._element == p45._element:
                    p45_idx = idx
                    break
            if p45_idx is not None and p45_idx >= 2:
                p_del1 = doc.paragraphs[p45_idx - 1]
                p_del2 = doc.paragraphs[p45_idx - 2]
                if p_del1.text.strip() == "" and p_del2.text.strip() == "":
                    p_del1._element.getparent().remove(p_del1._element)
                    p_del2._element.getparent().remove(p_del2._element)
        except Exception as e:
            print(f"Warning deleting spacing paragraphs: {e}")

    # 6. Update Table 0
    t = doc.tables[0]

    # Clear fixed row height of Row 1 so it shrinks and auto-adjusts to text height
    row1 = t.rows[1]
    if row1._tr.trPr is not None:
        trHeight = row1._tr.trPr.find(qn('w:trHeight'))
        if trHeight is not None:
            row1._tr.trPr.remove(trHeight)

    lignes = data.get("lignes", [])

    # Calculate totals
    total_ht = 0
    for idx, ligne in enumerate(lignes):
        nb_seaux = float(ligne.get("nbSeaux", 0))
        qty_seau_kg = float(ligne.get("qtySeauKg", 0))
        pu = float(ligne.get("prixUnitaire", 0))
        
        qty_kg = nb_seaux * qty_seau_kg
        montant_ht = qty_kg * pu
        total_ht += montant_ht
        
        # Display EMB as [nbSeaux]*[qtySeauKg]
        emb_display = f"{int(nb_seaux)}*{int(qty_seau_kg)}" if nb_seaux.is_integer() and qty_seau_kg.is_integer() else f"{nb_seaux}*{qty_seau_kg}"
        ligne["emb"] = emb_display
        ligne["qty_kg"] = qty_kg
        ligne["montant_ht"] = montant_ht

    total_tva = total_ht * 0.20
    total_ttc = total_ht + total_tva

    # Write first item to Row 1
    if len(lignes) > 0:
        first_item = lignes[0]
        set_cell_margins(t.rows[1].cells[0], left=100, right=150)
        set_cell_text(t.rows[1].cells[0], first_item["emb"])
        remove_cell_horizontal_borders(t.rows[1].cells[0])
        
        set_cell_margins(t.rows[1].cells[1], left=150, right=100)
        set_cell_text(t.rows[1].cells[1], first_item["designation"])
        remove_cell_horizontal_borders(t.rows[1].cells[1])
        
        set_cell_margins(t.rows[1].cells[2], left=100, right=100)
        set_cell_text(t.rows[1].cells[2], format_money(first_item["qty_kg"]))
        remove_cell_horizontal_borders(t.rows[1].cells[2])
        
        set_cell_margins(t.rows[1].cells[3], left=100, right=100)
        set_cell_text(t.rows[1].cells[3], format_money(first_item["prixUnitaire"]))
        remove_cell_horizontal_borders(t.rows[1].cells[3])
        
        # Remove margins and write total
        set_cell_margins(t.rows[1].cells[4], left=10, right=10)
        set_cell_text(t.rows[1].cells[4], format_money(first_item["montant_ht"]), right_align=True, font_size=12)
        remove_cell_horizontal_borders(t.rows[1].cells[4])

    # Write subsequent items by cloning Row 1
    for idx in range(1, len(lignes)):
        item = lignes[idx]
        row1 = t.rows[1]
        new_tr = copy.deepcopy(row1._tr)
        
        # Insert new row element before total row (originally index 2, now 2 + idx - 1 = 1 + idx)
        target_tr = t.rows[1 + idx]._tr
        parent = target_tr.getparent()
        parent.insert(parent.index(target_tr), new_tr)
        
        new_row = t.rows[1 + idx]
        set_cell_margins(new_row.cells[0], left=100, right=150)
        set_cell_text(new_row.cells[0], item["emb"])
        remove_cell_horizontal_borders(new_row.cells[0])
        
        set_cell_margins(new_row.cells[1], left=150, right=100)
        set_cell_text(new_row.cells[1], item["designation"])
        remove_cell_horizontal_borders(new_row.cells[1])
        
        set_cell_margins(new_row.cells[2], left=100, right=100)
        set_cell_text(new_row.cells[2], format_money(item["qty_kg"]))
        remove_cell_horizontal_borders(new_row.cells[2])
        
        set_cell_margins(new_row.cells[3], left=100, right=100)
        set_cell_text(new_row.cells[3], format_money(item["prixUnitaire"]))
        remove_cell_horizontal_borders(new_row.cells[3])
        
        # Remove margins and write total
        set_cell_margins(new_row.cells[4], left=10, right=10)
        set_cell_text(new_row.cells[4], format_money(item["montant_ht"]), right_align=True, font_size=12)
        remove_cell_horizontal_borders(new_row.cells[4])

    # Update totals rows (last 3 rows)
    num_rows = len(t.rows)
    row_total_ht = t.rows[num_rows - 3]
    row_total_tva = t.rows[num_rows - 2]
    row_total_ttc = t.rows[num_rows - 1]

    # Remove margins of total value cells to prevent wrapping at 14pt
    set_cell_margins(row_total_ht.cells[3], left=0, right=0)
    set_cell_margins(row_total_tva.cells[3], left=0, right=0)
    set_cell_margins(row_total_ttc.cells[3], left=0, right=0)

    set_cell_text(row_total_ht.cells[3], format_money(total_ht), right_align=True, font_size=14)
    set_cell_text(row_total_tva.cells[3], format_money(total_tva), right_align=True, font_size=14)
    set_cell_text(row_total_ttc.cells[3], format_money(total_ttc), right_align=True, font_size=14)

    doc.save(output_path)
    print("Docx generation successful!")

if __name__ == "__main__":
    main()
